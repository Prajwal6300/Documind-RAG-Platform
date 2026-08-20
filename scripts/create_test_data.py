"""Creates sample documents for the accuracy test suite."""

import os
import sys
from pathlib import Path

ROOT_DIR = Path(__file__).resolve().parent.parent
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

sys.stdout.reconfigure(encoding="utf-8")

TEST_DATA_DIR = ROOT_DIR / "test_data"


def create_company_policy_pdf(path=None):
    import pymupdf

    path = Path(path) if path else TEST_DATA_DIR / "company_policy.pdf"
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = pymupdf.open()
    doc.new_page(width=595, height=842)
    doc.new_page(width=595, height=842)
    doc.new_page(width=595, height=842)

    page1 = doc[0]
    page1.insert_text((72, 72), "Company Leave Policy 2025", fontsize=16)
    page1.insert_text((72, 100), "1. Casual Leave: Every employee is entitled to 12 days of casual leave per year.")
    page1.insert_text((72, 130), "2. Sick Leave: Every employee is entitled to 10 days of sick leave per year.")
    page1.insert_text((72, 160), "3. Annual Leave: After one year of employment, employees earn 20 days of annual leave.")

    page2 = doc[1]
    page2.insert_text((72, 72), "Working Hours", fontsize=16)
    page2.insert_text((72, 100), "Working hours are from 9 AM to 6 PM, Monday through Friday.")
    page2.insert_text((72, 130), "Lunch break is between 1 PM and 2 PM.")

    page3 = doc[2]
    page3.insert_text((72, 72), "Employee Benefits", fontsize=16)
    page3.insert_text((72, 100), "Employee ID: EMP1024")
    page3.insert_text((72, 130), "Health insurance covers up to $50,000 per year.")
    page3.insert_text((72, 160), "Joining date: 15 March 2021")
    page3.insert_text((72, 190), "Monthly salary: $4,500.00")

    doc.save(str(path))
    doc.close()
    print(f"Created {path}")


def create_employee_handbook_pdf(path=None):
    import pymupdf

    path = Path(path) if path else TEST_DATA_DIR / "employee_handbook.pdf"
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = pymupdf.open()
    doc.new_page(width=595, height=842)

    page1 = doc[0]
    page1.insert_text((72, 72), "Employee Handbook", fontsize=16)
    page1.insert_text((72, 100), "Code of conduct: All employees must follow the company code of conduct.")
    page1.insert_text((72, 130), "Workplace safety guidelines are provided by the safety department.")
    page1.insert_text((72, 160), "Communication: Use the official communication channels for work.")

    doc.save(str(path))
    doc.close()
    print(f"Created {path}")


def create_resume_txt(path=None):
    path = Path(path) if path else TEST_DATA_DIR / "resume.txt"
    path.parent.mkdir(parents=True, exist_ok=True)

    content = """Prajwal M. S. Yadav
Software Engineer

Skills:
- Python
- React
- Node.js
- SQL

Experience:
- Full Stack Developer at TechCorp (2020-2024)
- AI/ML projects with Python

Education:
- B.Tech Computer Science

Contact: +91-98765-43210
Email: prajwal.yadav@example.com
"""
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)

    print(f"Created {path}")


def create_inventory_docx(path=None):
    from docx import Document

    path = Path(path) if path else TEST_DATA_DIR / "inventory.docx"
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Product Inventory", 0)
    doc.add_paragraph("The following table lists all products and their prices.")

    table = doc.add_table(rows=4, cols=3)
    table.cell(0, 0).text = "Product"
    table.cell(0, 1).text = "Price"
    table.cell(0, 2).text = "Quantity"
    table.cell(1, 0).text = "Hex Bolt M12"
    table.cell(1, 1).text = "$1.50"
    table.cell(1, 2).text = "5000"
    table.cell(2, 0).text = "Steel Washer"
    table.cell(2, 1).text = "$0.20"
    table.cell(2, 2).text = "10000"
    table.cell(3, 0).text = "Aluminium Bracket"
    table.cell(3, 1).text = "$8.75"
    table.cell(3, 2).text = "2000"

    doc.save(str(path))
    print(f"Created {path}")


if __name__ == "__main__":
    create_company_policy_pdf()
    create_employee_handbook_pdf()
    create_resume_txt()
    create_inventory_docx()
